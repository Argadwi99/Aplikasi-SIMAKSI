import streamlit as st
import docx
from docx import Document
import io
from datetime import date
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# --- KONFIGURASI PENTING UNTUK EMAIL ---
# Cukup isi dua baris ini. Gunakan App Password, BUKAN password email utama.
ADMIN_EMAIL = "abksda8@gmail.com"      # GANTI DENGAN EMAIL ADMIN ANDA
APP_PASSWORD = "ijux pfhg izvq cjbf"       # GANTI DENGAN 16 KARAKTER APP PASSWORD ANDA

# ==============================================================================
# FUNGSI-FUNGSI PEMBANTU
# ==============================================================================
def fill_word_template(template_path, data_dict):
    """Mengisi template Word dengan data yang diberikan."""
    try:
        doc = Document(template_path)
        for p in doc.paragraphs:
            for key, value in data_dict.items():
                if key in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, value)
                            inline[i].text = text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in data_dict.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, value)
                                        inline[i].text = text
        return doc
    except Exception as e:
        st.error(f"Gagal memproses template: {e}")
        return None

def kirim_email(dokumen_word, nama_file, nama_pemohon):
    """Mengirim email dengan file Word sebagai lampiran."""
    try:
        msg = MIMEMultipart()
        msg['From'] = ADMIN_EMAIL
        msg['To'] = ADMIN_EMAIL # Mengirim ke alamat email yang sama
        msg['Subject'] = f"Dokumen SIMAKSI Baru - a/n {nama_pemohon}"

        body = f"Dokumen SIMAKSI baru telah dibuat atas nama {nama_pemohon}.\n\nFile terlampir."
        msg.attach(MIMEText(body, 'plain'))

        lampiran = MIMEApplication(dokumen_word, Name=nama_file)
        lampiran['Content-Disposition'] = f'attachment; filename="{nama_file}"'
        msg.attach(lampiran)

        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(ADMIN_EMAIL, APP_PASSWORD)
            server.sendmail(ADMIN_EMAIL, ADMIN_EMAIL, msg.as_string())
        return True
    except Exception as e:
        st.error(f"Gagal mengirim email: {e}")
        return False

# ==============================================================================
# TAMPILAN APLIKASI WEB
# ==============================================================================
st.title("Aplikasi Pembuat Dokumen SIMAKSI")
st.write("Silakan isi form di bawah ini. Dokumen akan otomatis terkirim ke email administrasi.")

with st.form("simaksi_form"):
    nama_lengkap = st.text_input("Nama Lengkap:")
    tujuan_kegiatan = st.text_input("Tujuan Kegiatan:")
    lokasi_konservasi = st.text_input("Lokasi Kawasan Konservasi:")
    tanggal_mulai_obj = st.date_input("Tanggal Mulai:", value=date.today(), format="DD/MM/YYYY")
    tanggal_selesai_obj = st.date_input("Tanggal Selesai:", value=date.today(), format="DD/MM/YYYY")
    jumlah_peserta = st.text_input("Jumlah Peserta/Pengikut:")
    
    submitted = st.form_submit_button("Buat & Kirim Dokumen ke Admin")

if submitted:
    tanggal_mulai = tanggal_mulai_obj.strftime("%d/%m/%Y")
    tanggal_selesai = tanggal_selesai_obj.strftime("%d/%m/%Y")

    if not all([nama_lengkap, tujuan_kegiatan, lokasi_konservasi, tanggal_mulai, tanggal_selesai, jumlah_peserta]):
        st.warning("Harap isi semua kolom.")
    else:
        with st.spinner("Sedang memproses dan mengirim dokumen..."):
            data_to_fill = {
                '[nama_lengkap]': nama_lengkap,
                '[tujuan_kegiatan]': tujuan_kegiatan,
                '[lokasi_konservasi]': lokasi_konservasi,
                '[tanggal_mulai]': tanggal_mulai,
                '[tanggal_selesai]': tanggal_selesai,
                '[jumlah_peserta]': jumlah_peserta,
            }
            
            template_name = "Draft SIMAKSI_Kosong.docx"
            filled_doc = fill_word_template(template_name, data_to_fill)

            if filled_doc:
                bio = io.BytesIO()
                filled_doc.save(bio)
                nama_file_output = f"SIMAKSI_{nama_lengkap.replace(' ', '_')}.docx"
                
                email_sent = kirim_email(bio.getvalue(), nama_file_output, nama_lengkap)
                
                if email_sent:
                    st.success("Dokumen berhasil dibuat dan telah terkirim ke email admin!")